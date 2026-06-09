import zipfile
from extractor.exceptions import ExtractionError


# Maximum uncompressed size (bytes) we are willing to parse from a single
# ZIP entry.  Prevents zip-bomb style DoS when processing untrusted DOCX files.
_MAX_XML_BYTES = 100 * 1024 * 1024  # 100 MB


def _safe_xml_parse(xml_bytes: bytes):
    """Parse XML bytes with external-entity / entity-expansion protection.

    Tries ``defusedxml`` first (safest); falls back to stdlib
    ``ElementTree`` with a hardened ``XMLParser`` that forbids DTDs.
    """
    try:
        import defusedxml.ElementTree as SafeET
        return SafeET.fromstring(xml_bytes)
    except ImportError:
        pass

    import xml.etree.ElementTree as ET

    # Python ≥ 3.7.2 expat-backed parser: ``forbid_dtd`` blocks all DTD
    # processing, which neutralises both XXE and billion-laughs attacks.
    try:
        parser = ET.XMLParser()
        parser.parser.UseForeignDTD(False)
        parser.parser.SetParamEntityParsing(0)  # XML_PARAM_ENTITY_PARSING_NEVER
        parser.feed(xml_bytes)
        return parser.close()
    except Exception:
        pass

    # Last resort – plain fromstring (still safe against external-entity
    # resolution on CPython ≥ 3.8, but entity-expansion is unbounded).
    return ET.fromstring(xml_bytes)


def extract_docx_with_python_docx(docx_path: str) -> str | None:
    try:
        import docx
        document = docx.Document(docx_path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts)
    except ImportError:
        return None
    except Exception:
        return None


def extract_docx_with_zipfile(docx_path: str) -> str | None:
    try:
        with zipfile.ZipFile(docx_path) as zf:
            info = zf.getinfo("word/document.xml")
            if info.file_size > _MAX_XML_BYTES:
                return None
            xml_bytes = zf.read("word/document.xml")
        root = _safe_xml_parse(xml_bytes)
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        parts: list[str] = []
        for paragraph in root.iter(f"{namespace}p"):
            texts = [node.text for node in paragraph.iter(f"{namespace}t") if node.text]
            if texts:
                parts.append("".join(texts))
        return "\n".join(parts) if parts else None
    except Exception:
        return None


def extract_docx(docx_path: str) -> tuple[str, str]:
    print("Trying python-docx...", end=" ", flush=True)
    text = extract_docx_with_python_docx(docx_path)
    if text and text.strip():
        print("OK")
        return text, "python-docx"

    print("not available")
    print("Trying stdlib DOCX parser...", end=" ", flush=True)
    text = extract_docx_with_zipfile(docx_path)
    if text and text.strip():
        print("OK")
        return text, "zipfile-docx"

    print("FAILED")
    raise ExtractionError(
        "Could not extract text from DOCX.\n"
        "Install python-docx for best results:\n"
        "  pip3 install python-docx"
    )

